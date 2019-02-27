import random
import string
import time
import uuid
from docx import Document
import os
import sys

start_time = time.time()

# Iterate over the data and write it out row by row.
sheet_size = 5
need_size = int(sys.argv[1])
# need_size = 1024 * 1024 * 150

filename = str(uuid.uuid4()) + ".docx"


def generate_docx(m_count):
    blank_doc = Document()
    inserting = blank_doc.add_paragraph('python_confirmed')
    for i in range(0, m_count):
        inserting.add_run(id_generator())
    blank_doc.save(filename)

    size = os.stat(filename).st_size
    return size


# need some optimizations here
def id_generator(size=need_size // 42, chars=string.ascii_uppercase + string.digits):
    return ''.join(random.choice(chars) for _ in range(size))


count = 50
size50 = generate_docx(count)
k50 = -(-size50 // 50)
count = int(need_size / k50) + 1


size = generate_docx(count)
if size < need_size:
    count = int(count - (-(need_size - size) // k50))
    size = generate_docx(count)

# print(str(size / 1024 / 1024) + " MBytes")
# print("--- %s seconds ---" % (time.time() - start_time))
print(filename)
